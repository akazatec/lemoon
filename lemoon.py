#!/usr/bin/env python3

## Lemoon - Multi Tool
## Author: AKAZA SENZO

import os
import sys
import re
import platform
import socket
import uuid
import requests
import subprocess

try:
    from rich.console import Console
    from rich.panel import Panel
    from rich.table import Table
    from rich import box
except ImportError:
    os.system("pip3 install rich")
    from rich.console import Console
    from rich.panel import Panel
    from rich.table import Table
    from rich import box

try:
    from pyfiglet import figlet_format
except ImportError:
    os.system("pip3 install pyfiglet")
    from pyfiglet import figlet_format

try:
    from colorama import init
    init(autoreset=True)
except ImportError:
    os.system("pip3 install colorama")
    from colorama import init
    init(autoreset=True)

try:
    from PIL import Image
    from PIL.ExifTags import TAGS, GPSTAGS
except ImportError:
    os.system("pip3 install Pillow")
    from PIL import Image
    from PIL.ExifTags import TAGS, GPSTAGS

try:
    from mutagen.mp3 import MP3
    from mutagen.mp4 import MP4
    from mutagen.flac import FLAC
    from mutagen.oggvorbis import OggVorbis
except ImportError:
    os.system("pip3 install mutagen")
    from mutagen.mp3 import MP3
    from mutagen.mp4 import MP4
    from mutagen.flac import FLAC
    from mutagen.oggvorbis import OggVorbis

try:
    import PyPDF2
except ImportError:
    os.system("pip3 install PyPDF2")
    import PyPDF2

try:
    import docx
    from docx import Document
except ImportError:
    os.system("pip3 install python-docx")
    import docx
    from docx import Document

try:
    from bs4 import BeautifulSoup
except ImportError:
    os.system("pip3 install beautifulsoup4")
    from bs4 import BeautifulSoup

console = Console()

def clear():
    os.system('clear')

def banner():
    clear()
    fig = figlet_format("Lemoon", font="doom")
    console.print(fig, style="bold yellow")
    console.print(Panel.fit(
        "[bold yellow]TOOL   :[/] [white]Lemoon - Multi Tool Suite[/]\n"
        "[bold yellow]AUTHOR :[/] [white]AKAZA SENZO[/]\n"
        "[bold yellow]VERSION:[/] [white]2.0.0[/]\n"
        "[bold yellow][ FOR EDUCATIONAL PURPOSES ONLY ][/]",
        border_style="yellow"
    ))
    print()

# ══════════════════════════════════════
# 1. METAEXTRACT
# ══════════════════════════════════════

def get_gps(exif_data):
    gps_info = {}
    if "GPSInfo" in exif_data:
        for key in exif_data["GPSInfo"].keys():
            decode = GPSTAGS.get(key, key)
            gps_info[decode] = exif_data["GPSInfo"][key]
    return gps_info

def convert_gps(gps_info):
    try:
        lat = gps_info.get("GPSLatitude")
        lat_ref = gps_info.get("GPSLatitudeRef")
        lon = gps_info.get("GPSLongitude")
        lon_ref = gps_info.get("GPSLongitudeRef")
        if lat and lon:
            lat = float(lat[0]) + float(lat[1])/60 + float(lat[2])/3600
            lon = float(lon[0]) + float(lon[1])/60 + float(lon[2])/3600
            if lat_ref == "S": lat = -lat
            if lon_ref == "W": lon = -lon
            return f"{lat:.6f}, {lon:.6f}"
    except:
        pass
    return None

def extract_image(filepath):
    banner()
    table = Table(title=f"[bold yellow]Image Metadata: {os.path.basename(filepath)}[/]",
                 box=box.DOUBLE_EDGE, border_style="yellow")
    table.add_column("Field", style="yellow", no_wrap=True)
    table.add_column("Value", style="white")
    try:
        img = Image.open(filepath)
        table.add_row("File", os.path.basename(filepath))
        table.add_row("Format", str(img.format))
        table.add_row("Mode", str(img.mode))
        table.add_row("Size", f"{img.size[0]}x{img.size[1]} px")
        table.add_row("File Size", f"{os.path.getsize(filepath)} bytes")
        exif_data = {}
        raw = img._getexif()
        if raw:
            for tag_id, value in raw.items():
                tag = TAGS.get(tag_id, tag_id)
                exif_data[tag] = value
            important = [
                "Make", "Model", "Software", "DateTime",
                "DateTimeOriginal", "DateTimeDigitized",
                "ExifImageWidth", "ExifImageHeight",
                "Flash", "FocalLength", "ISOSpeedRatings",
                "ExposureTime", "FNumber", "WhiteBalance",
                "LightSource", "Artist", "Copyright",
                "ImageDescription", "Orientation"
            ]
            for field in important:
                if field in exif_data:
                    table.add_row(field, str(exif_data[field]))
            gps = get_gps(exif_data)
            if gps:
                coords = convert_gps(gps)
                if coords:
                    table.add_row("GPS Coordinates", coords)
                    table.add_row("Google Maps", f"https://maps.google.com/?q={coords}")
        else:
            table.add_row("EXIF Data", "No EXIF data found")
    except Exception as e:
        table.add_row("Error", str(e))
    console.print(table)

def extract_audio(filepath):
    banner()
    table = Table(title=f"[bold yellow]Audio Metadata: {os.path.basename(filepath)}[/]",
                 box=box.DOUBLE_EDGE, border_style="yellow")
    table.add_column("Field", style="yellow", no_wrap=True)
    table.add_column("Value", style="white")
    try:
        ext = filepath.lower().split(".")[-1]
        table.add_row("File", os.path.basename(filepath))
        table.add_row("File Size", f"{os.path.getsize(filepath)} bytes")
        if ext == "mp3":
            audio = MP3(filepath)
            table.add_row("Duration", f"{audio.info.length:.2f} sec")
            table.add_row("Bitrate", f"{audio.info.bitrate} bps")
            table.add_row("Sample Rate", f"{audio.info.sample_rate} Hz")
            if audio.tags:
                for tag, val in audio.tags.items():
                    table.add_row(str(tag), str(val))
        elif ext in ["mp4", "m4a"]:
            audio = MP4(filepath)
            table.add_row("Duration", f"{audio.info.length:.2f} sec")
            tag_map = {
                "\xa9nam": "Title", "\xa9ART": "Artist",
                "\xa9alb": "Album", "\xa9day": "Year",
                "\xa9gen": "Genre", "cprt": "Copyright"
            }
            for tag, name in tag_map.items():
                if tag in audio:
                    table.add_row(name, str(audio[tag][0]))
        elif ext == "flac":
            audio = FLAC(filepath)
            table.add_row("Duration", f"{audio.info.length:.2f} sec")
            if audio.tags:
                for tag, val in audio.tags.items():
                    table.add_row(str(tag), str(val[0]))
        elif ext == "ogg":
            audio = OggVorbis(filepath)
            table.add_row("Duration", f"{audio.info.length:.2f} sec")
            if audio.tags:
                for tag, val in audio.tags.items():
                    table.add_row(str(tag), str(val[0]))
    except Exception as e:
        table.add_row("Error", str(e))
    console.print(table)

def extract_video(filepath):
    banner()
    table = Table(title=f"[bold yellow]Video Metadata: {os.path.basename(filepath)}[/]",
                 box=box.DOUBLE_EDGE, border_style="yellow")
    table.add_column("Field", style="yellow", no_wrap=True)
    table.add_column("Value", style="white")
    try:
        table.add_row("File", os.path.basename(filepath))
        table.add_row("File Size", f"{os.path.getsize(filepath)} bytes")
        video = MP4(filepath)
        table.add_row("Duration", f"{video.info.length:.2f} sec")
        tag_map = {
            "\xa9nam": "Title", "\xa9ART": "Artist",
            "\xa9alb": "Album", "\xa9day": "Year",
            "\xa9gen": "Genre", "cprt": "Copyright"
        }
        for tag, name in tag_map.items():
            if tag in video:
                table.add_row(name, str(video[tag][0]))
    except Exception as e:
        table.add_row("Error", str(e))
    console.print(table)

def extract_pdf(filepath):
    banner()
    table = Table(title=f"[bold yellow]PDF Metadata: {os.path.basename(filepath)}[/]",
                 box=box.DOUBLE_EDGE, border_style="yellow")
    table.add_column("Field", style="yellow", no_wrap=True)
    table.add_column("Value", style="white")
    try:
        table.add_row("File", os.path.basename(filepath))
        table.add_row("File Size", f"{os.path.getsize(filepath)} bytes")
        with open(filepath, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            table.add_row("Pages", str(len(reader.pages)))
            table.add_row("Encrypted", str(reader.is_encrypted))
            meta = reader.metadata
            if meta:
                for field in ["/Title","/Author","/Subject","/Creator",
                              "/Producer","/CreationDate","/ModDate","/Keywords"]:
                    val = meta.get(field)
                    if val:
                        table.add_row(field.replace("/",""), str(val))
    except Exception as e:
        table.add_row("Error", str(e))
    console.print(table)

def extract_docx_file(filepath):
    banner()
    table = Table(title=f"[bold yellow]DOCX Metadata: {os.path.basename(filepath)}[/]",
                 box=box.DOUBLE_EDGE, border_style="yellow")
    table.add_column("Field", style="yellow", no_wrap=True)
    table.add_column("Value", style="white")
    try:
        table.add_row("File", os.path.basename(filepath))
        table.add_row("File Size", f"{os.path.getsize(filepath)} bytes")
        doc = Document(filepath)
        props = doc.core_properties
        table.add_row("Title", str(props.title) if props.title else "N/A")
        table.add_row("Author", str(props.author) if props.author else "N/A")
        table.add_row("Last Modified By", str(props.last_modified_by) if props.last_modified_by else "N/A")
        table.add_row("Created", str(props.created) if props.created else "N/A")
        table.add_row("Modified", str(props.modified) if props.modified else "N/A")
        table.add_row("Description", str(props.description) if props.description else "N/A")
        table.add_row("Keywords", str(props.keywords) if props.keywords else "N/A")
        table.add_row("Revision", str(props.revision) if props.revision else "N/A")
        table.add_row("Paragraphs", str(len(doc.paragraphs)))
    except Exception as e:
        table.add_row("Error", str(e))
    console.print(table)

def auto_extract(filepath):
    ext = filepath.lower().split(".")[-1]
    if ext in ["jpg","jpeg","png","tiff","bmp","gif","webp"]:
        extract_image(filepath)
    elif ext in ["mp3","flac","ogg","m4a"]:
        extract_audio(filepath)
    elif ext in ["mp4","m4v","mov"]:
        extract_video(filepath)
    elif ext == "pdf":
        extract_pdf(filepath)
    elif ext == "docx":
        extract_docx_file(filepath)
    else:
        console.print(f"[red][!] Unsupported: .{ext}[/]")

def metaextract_menu():
    while True:
        banner()
        console.print(Panel(
            "[bold yellow][1][/]  Image Metadata  (JPG,PNG,TIFF,BMP,GIF,WEBP)\n"
            "[bold yellow][2][/]  Audio Metadata  (MP3,FLAC,OGG,M4A)\n"
            "[bold yellow][3][/]  Video Metadata  (MP4,MOV,M4V)\n"
            "[bold yellow][4][/]  PDF Metadata\n"
            "[bold yellow][5][/]  DOCX Metadata\n"
            "[bold yellow][6][/]  Auto Detect & Extract\n"
            "[bold red][0][/]   Back",
            title="[bold yellow][ MetaExtract ][/]",
            border_style="yellow"
        ))
        console.print("[bold yellow]┌──[[white]AKAZA SENZO[yellow]㉿[white]Lemoon[yellow]]-[[white]~[yellow]][/]")
        choice = console.input("[bold yellow]└─[white]$ [/]").strip()
        if choice == "0":
            break
        elif choice in ["1","2","3","4","5","6"]:
            filepath = console.input("[bold yellow]└─[white]$ File Path: [/]").strip().strip("'\"")
            if not os.path.exists(filepath):
                console.print("[red][!] File not found![/]")
            else:
                if choice == "1": extract_image(filepath)
                elif choice == "2": extract_audio(filepath)
                elif choice == "3": extract_video(filepath)
                elif choice == "4": extract_pdf(filepath)
                elif choice == "5": extract_docx_file(filepath)
                elif choice == "6": auto_extract(filepath)
        else:
            console.print("[red][!] Invalid![/]")
        console.input("\n[bold yellow][Press Enter...][/]")

# ══════════════════════════════════════
# 2. META EDITOR
# ══════════════════════════════════════

def edit_image_meta(filepath):
    banner()
    console.print("[bold yellow][!] Removing all EXIF data (privacy clean)[/]\n")
    try:
        img = Image.open(filepath)
        data = list(img.getdata())
        clean_img = Image.new(img.mode, img.size)
        clean_img.putdata(data)
        out = filepath.rsplit(".", 1)[0] + "_cleaned." + filepath.rsplit(".", 1)[1]
        clean_img.save(out)
        console.print(f"[bold green][✓] Cleaned image saved: {out}[/]")
    except Exception as e:
        console.print(f"[red][!] Error: {e}[/]")

def edit_docx_meta(filepath):
    banner()
    console.print("[bold yellow][?] Enter new metadata values (leave blank to skip)[/]\n")
    try:
        doc = Document(filepath)
        props = doc.core_properties
        title = console.input("[bold yellow]└─[white]$ New Title: [/]").strip()
        author = console.input("[bold yellow]└─[white]$ New Author: [/]").strip()
        subject = console.input("[bold yellow]└─[white]$ New Subject: [/]").strip()
        keywords = console.input("[bold yellow]└─[white]$ New Keywords: [/]").strip()
        description = console.input("[bold yellow]└─[white]$ New Description: [/]").strip()
        if title: props.title = title
        if author: props.author = author
        if subject: props.subject = subject
        if keywords: props.keywords = keywords
        if description: props.description = description
        out = filepath.rsplit(".", 1)[0] + "_edited.docx"
        doc.save(out)
        console.print(f"[bold green][✓] Edited DOCX saved: {out}[/]")
    except Exception as e:
        console.print(f"[red][!] Error: {e}[/]")

def edit_audio_meta(filepath):
    banner()
    console.print("[bold yellow][?] Enter new metadata values (leave blank to skip)[/]\n")
    try:
        ext = filepath.lower().split(".")[-1]
        if ext == "mp3":
            from mutagen.easyid3 import EasyID3
            audio = EasyID3(filepath)
            title = console.input("[bold yellow]└─[white]$ New Title: [/]").strip()
            artist = console.input("[bold yellow]└─[white]$ New Artist: [/]").strip()
            album = console.input("[bold yellow]└─[white]$ New Album: [/]").strip()
            year = console.input("[bold yellow]└─[white]$ New Year: [/]").strip()
            genre = console.input("[bold yellow]└─[white]$ New Genre: [/]").strip()
            if title: audio["title"] = title
            if artist: audio["artist"] = artist
            if album: audio["album"] = album
            if year: audio["date"] = year
            if genre: audio["genre"] = genre
            audio.save()
            console.print(f"[bold green][✓] MP3 metadata updated![/]")
        elif ext in ["mp4","m4a"]:
            audio = MP4(filepath)
            title = console.input("[bold yellow]└─[white]$ New Title: [/]").strip()
            artist = console.input("[bold yellow]└─[white]$ New Artist: [/]").strip()
            album = console.input("[bold yellow]└─[white]$ New Album: [/]").strip()
            year = console.input("[bold yellow]└─[white]$ New Year: [/]").strip()
            if title: audio["\xa9nam"] = [title]
            if artist: audio["\xa9ART"] = [artist]
            if album: audio["\xa9alb"] = [album]
            if year: audio["\xa9day"] = [year]
            audio.save()
            console.print(f"[bold green][✓] MP4 metadata updated![/]")
        elif ext == "flac":
            audio = FLAC(filepath)
            title = console.input("[bold yellow]└─[white]$ New Title: [/]").strip()
            artist = console.input("[bold yellow]└─[white]$ New Artist: [/]").strip()
            album = console.input("[bold yellow]└─[white]$ New Album: [/]").strip()
            if title: audio["title"] = title
            if artist: audio["artist"] = artist
            if album: audio["album"] = album
            audio.save()
            console.print(f"[bold green][✓] FLAC metadata updated![/]")
        else:
            console.print("[red][!] Unsupported format[/]")
    except Exception as e:
        console.print(f"[red][!] Error: {e}[/]")

def metaeditor_menu():
    while True:
        banner()
        console.print(Panel(
            "[bold yellow][1][/]  Edit Image Metadata  (Remove EXIF)\n"
            "[bold yellow][2][/]  Edit Audio Metadata  (MP3,FLAC,M4A)\n"
            "[bold yellow][3][/]  Edit DOCX Metadata\n"
            "[bold red][0][/]   Back",
            title="[bold yellow][ Meta Editor ][/]",
            border_style="yellow"
        ))
        console.print("[bold yellow]┌──[[white]AKAZA SENZO[yellow]㉿[white]Lemoon[yellow]]-[[white]~[yellow]][/]")
        choice = console.input("[bold yellow]└─[white]$ [/]").strip()
        if choice == "0":
            break
        elif choice in ["1","2","3"]:
            filepath = console.input("[bold yellow]└─[white]$ File Path: [/]").strip().strip("'\"")
            if not os.path.exists(filepath):
                console.print("[red][!] File not found![/]")
            else:
                if choice == "1": edit_image_meta(filepath)
                elif choice == "2": edit_audio_meta(filepath)
                elif choice == "3": edit_docx_meta(filepath)
        else:
            console.print("[red][!] Invalid![/]")
        console.input("\n[bold yellow][Press Enter...][/]")

# ══════════════════════════════════════
# 3. DORKHUNTER
# ══════════════════════════════════════

DORKS = {
    "Login Pages": 'inurl:login',
    "Admin Panels": 'inurl:admin',
    "Config Files": 'ext:cfg OR ext:conf OR ext:config',
    "Backup Files": 'ext:bak OR ext:backup OR ext:old',
    "Database Files": 'ext:sql OR ext:db OR ext:sqlite',
    "Log Files": 'ext:log',
    "Password Files": 'inurl:password OR inurl:passwd',
    "PHP Files": 'ext:php inurl:index',
    "XML Files": 'ext:xml',
    "Excel Files": 'ext:xls OR ext:xlsx',
    "PDF Documents": 'ext:pdf',
    "Open Directories": 'intitle:"index of"',
    "WordPress Login": 'inurl:wp-login.php',
    "phpMyAdmin": 'inurl:phpmyadmin',
    "FTP Login": 'inurl:ftp',
    "Camera Feeds": 'inurl:view/index.shtml',
    "Vulnerable Forms": 'inurl:search.php?q=',
    "Error Messages": 'intext:"sql syntax near"',
    "Sensitive Docs": 'intitle:"confidential" ext:doc OR ext:pdf',
    "Email Lists": 'ext:txt inurl:email OR inurl:emails',
}

def dorkhunter_menu():
    while True:
        banner()
        console.print(Panel(
            "[bold yellow][1][/]  Predefined Dorks List\n"
            "[bold yellow][2][/]  Custom Dork + Target\n"
            "[bold yellow][3][/]  Generate Dork URLs\n"
            "[bold red][0][/]   Back",
            title="[bold yellow][ DorkHunter ][/]",
            border_style="yellow"
        ))
        console.print("[bold yellow]┌──[[white]AKAZA SENZO[yellow]㉿[white]Lemoon[yellow]]-[[white]~[yellow]][/]")
        choice = console.input("[bold yellow]└─[white]$ [/]").strip()

        if choice == "0":
            break
        elif choice == "1":
            banner()
            console.print("[bold yellow][*] Predefined Dorks:[/]\n")
            for i, (cat, dork) in enumerate(DORKS.items(), 1):
                console.print(f"[bold yellow][{i}] {cat}[/]")
                print(dork)
                print()
        elif choice == "2":
            banner()
            target = console.input("[bold yellow]└─[white]$ Enter Target Domain: [/]").strip()
            console.print(f"\n[bold yellow][*] Dorks for: {target}[/]\n")
            for cat, dork in DORKS.items():
                url = f"https://www.google.com/search?q=site:{target}+{dork.replace(' ', '+')}"
                console.print(f"[bold yellow][{cat}][/]")
                print(url)
                print()
        elif choice == "3":
            banner()
            custom = console.input("[bold yellow]└─[white]$ Enter Custom Dork: [/]").strip()
            target = console.input("[bold yellow]└─[white]$ Target (optional, Enter to skip): [/]").strip()
            if target:
                url = f"https://www.google.com/search?q=site:{target}+{custom.replace(' ', '+')}"
            else:
                url = f"https://www.google.com/search?q={custom.replace(' ', '+')}"
            console.print(f"\n[bold yellow][Dork][/] {custom}")
            console.print(f"[bold yellow][Target][/] {target if target else 'None'}")
            console.print(f"[bold yellow][URL][/]")
            print(url)
        else:
            console.print("[red][!] Invalid![/]")
        console.input("\n[bold yellow][Press Enter...][/]")

# ══════════════════════════════════════
# 4. IMAGE REVERSE
# ══════════════════════════════════════

def imgreverse_menu():
    while True:
        banner()
        console.print(Panel(
            "[bold yellow][1][/]  Google Reverse Image Search\n"
            "[bold yellow][2][/]  Yandex Reverse Image Search\n"
            "[bold yellow][3][/]  Bing Reverse Image Search\n"
            "[bold yellow][4][/]  TinEye Reverse Image Search\n"
            "[bold yellow][5][/]  All Search Engines\n"
            "[bold red][0][/]   Back",
            title="[bold yellow][ Image Reverse ][/]",
            border_style="yellow"
        ))
        console.print("[bold yellow]┌──[[white]AKAZA SENZO[yellow]㉿[white]Lemoon[yellow]]-[[white]~[yellow]][/]")
        choice = console.input("[bold yellow]└─[white]$ [/]").strip()

        if choice == "0":
            break

        if choice in ["1","2","3","4","5"]:
            img_input = console.input("[bold yellow]└─[white]$ Enter Image URL: [/]").strip().strip("'\"")
            from urllib.parse import quote
            encoded_url = quote(img_input, safe='')
            google_url = f"https://lens.google.com/uploadbyurl?url={encoded_url}"
            yandex_url = f"https://yandex.com/images/search?url={encoded_url}&rpt=imageview"
            bing_url = f"https://www.bing.com/images/search?q=imgurl:{encoded_url}&view=detailv2&iss=sbi"
            tineye_url = f"https://tineye.com/search?url={encoded_url}"

            console.print("\n[bold yellow][*] Search URLs:[/]\n")
            if choice == "1" or choice == "5":
                console.print("[bold yellow][Google Lens][/]")
                print(google_url)
                print()
            if choice == "2" or choice == "5":
                console.print("[bold yellow][Yandex][/]")
                print(yandex_url)
                print()
            if choice == "3" or choice == "5":
                console.print("[bold yellow][Bing][/]")
                print(bing_url)
                print()
            if choice == "4" or choice == "5":
                console.print("[bold yellow][TinEye][/]")
                print(tineye_url)
                print()
            console.print("[bold yellow][!] Copy URL and open in browser![/]")
        else:
            console.print("[red][!] Invalid![/]")
        console.input("\n[bold yellow][Press Enter...][/]")

# ══════════════════════════════════════
# 5. TERMFETCH
# ══════════════════════════════════════

def termfetch():
    banner()
    try:
        table = Table(title="[bold yellow]System Information[/]", box=box.DOUBLE_EDGE, border_style="yellow")
        table.add_column("Field", style="yellow", no_wrap=True)
        table.add_column("Value", style="white")
        table.add_row("OS", platform.system())
        table.add_row("OS Version", platform.version()[:50])
        table.add_row("OS Release", platform.release())
        table.add_row("Architecture", platform.machine())
        table.add_row("Processor", platform.processor()[:50] if platform.processor() else "N/A")
        table.add_row("Hostname", socket.gethostname())
        table.add_row("Python", platform.python_version())
        table.add_row("Node", platform.node())
        try:
            local_ip = socket.gethostbyname(socket.gethostname())
            table.add_row("Local IP", local_ip)
        except:
            table.add_row("Local IP", "N/A")
        try:
            mac = ':'.join(['{:02x}'.format((uuid.getnode() >> i) & 0xff)
                           for i in range(0,8*6,8)][::-1])
            table.add_row("MAC Address", mac)
        except:
            table.add_row("MAC Address", "N/A")
        try:
            pub = requests.get("http://ip-api.com/json/", timeout=5).json()
            table.add_row("Public IP", pub.get("query", "N/A"))
            table.add_row("Country", pub.get("country", "N/A"))
            table.add_row("City", pub.get("city", "N/A"))
            table.add_row("ISP", pub.get("isp", "N/A"))
            table.add_row("Timezone", pub.get("timezone", "N/A"))
        except:
            table.add_row("Public IP", "Could not fetch")
        try:
            with open("/proc/cpuinfo") as f:
                cpu_info = f.read()
            cpu_lines = [l for l in cpu_info.split("\n") if "model name" in l]
            if cpu_lines:
                table.add_row("CPU", cpu_lines[0].split(":")[1].strip()[:50])
            with open("/proc/meminfo") as f:
                mem_info = f.read()
            mem_lines = {l.split(":")[0]: l.split(":")[1].strip()
                        for l in mem_info.split("\n") if ":" in l}
            table.add_row("RAM Total", mem_lines.get("MemTotal", "N/A"))
            table.add_row("RAM Free", mem_lines.get("MemAvailable", "N/A"))
        except:
            table.add_row("CPU/RAM", "N/A")
        try:
            with open("/proc/uptime") as f:
                uptime_sec = float(f.read().split()[0])
            hours = int(uptime_sec // 3600)
            minutes = int((uptime_sec % 3600) // 60)
            table.add_row("Uptime", f"{hours}h {minutes}m")
        except:
            table.add_row("Uptime", "N/A")
        table.add_row("Shell", os.environ.get("SHELL", "N/A"))
        table.add_row("Terminal", os.environ.get("TERM", "N/A"))
        table.add_row("User", os.environ.get("USER", os.environ.get("USERNAME", "N/A")))
        console.print(table)
    except Exception as e:
        console.print(f"[red][!] Error: {e}[/]")

# ══════════════════════════════════════
# 6. ARCHIVEHUNT
# ══════════════════════════════════════

def archivehunt_menu():
    while True:
        banner()
        console.print(Panel(
            "[bold yellow][1][/]  URL History       — Wayback Machine snapshots\n"
            "[bold yellow][2][/]  Keyword Search    — Archive.org search\n"
            "[bold yellow][3][/]  Domain History    — Domain purane versions\n"
            "[bold yellow][4][/]  Save Page         — URL archive mein save karo\n"
            "[bold red][0][/]   Back",
            title="[bold yellow][ ArchiveHunt ][/]",
            border_style="yellow"
        ))
        console.print("[bold yellow]┌──[[white]AKAZA SENZO[yellow]㉿[white]Lemoon[yellow]]-[[white]~[yellow]][/]")
        choice = console.input("[bold yellow]└─[white]$ [/]").strip()

        if choice == "0":
            break

        elif choice == "1":
            banner()
            url = console.input("[bold yellow]└─[white]$ Enter URL: [/]").strip()
            if not url.startswith("http"):
                url = "https://" + url
            try:
                console.print(f"\n[bold yellow][*] Fetching snapshots for: {url}[/]\n")
                r = requests.get(
                    f"http://archive.org/wayback/available?url={url}",
                    timeout=10
                ).json()
                table = Table(title=f"[bold yellow]Wayback Snapshots: {url}[/]",
                            box=box.DOUBLE_EDGE, border_style="yellow")
                table.add_column("Field", style="yellow", no_wrap=True)
                table.add_column("Value", style="white")
                snap = r.get("archived_snapshots", {}).get("closest", {})
                if snap:
                    table.add_row("Status", str(snap.get("status", "N/A")))
                    table.add_row("Available", str(snap.get("available", "N/A")))
                    table.add_row("Timestamp", str(snap.get("timestamp", "N/A")))
                    table.add_row("Snapshot URL", "")
                    console.print(table)
                    print(snap.get("url", "N/A"))

                    # Get more snapshots
                    clean_url = url.replace("https://","").replace("http://","")
                    console.print(f"\n[bold yellow][*] All Snapshots URL:[/]")
                    print(f"https://web.archive.org/web/*/{clean_url}")
                    console.print(f"\n[bold yellow][*] CDX API — Full History:[/]")
                    print(f"http://web.archive.org/cdx/search/cdx?url={clean_url}&output=json&limit=10")
                else:
                    table.add_row("Result", "No snapshots found")
                    console.print(table)
            except Exception as e:
                console.print(f"[red][!] Error: {e}[/]")

        elif choice == "2":
            banner()
            keyword = console.input("[bold yellow]└─[white]$ Enter Keyword: [/]").strip()
            try:
                console.print(f"\n[bold yellow][*] Searching Archive.org for: {keyword}[/]\n")
                r = requests.get(
                    f"https://archive.org/advancedsearch.php?q={keyword.replace(' ', '+')}&output=json&rows=10",
                    timeout=10
                ).json()
                table = Table(title=f"[bold yellow]Archive Search: {keyword}[/]",
                            box=box.DOUBLE_EDGE, border_style="yellow")
                table.add_column("Title", style="yellow", no_wrap=True)
                table.add_column("Identifier", style="white")
                table.add_column("Year", style="yellow")
                docs = r.get("response", {}).get("docs", [])
                if docs:
                    for doc in docs:
                        table.add_row(
                            str(doc.get("title", "N/A"))[:40],
                            str(doc.get("identifier", "N/A")),
                            str(doc.get("year", "N/A"))
                        )
                    console.print(table)
                    console.print("\n[bold yellow][*] Archive URLs:[/]")
                    for doc in docs:
                        ident = doc.get("identifier", "")
                        if ident:
                            print(f"https://archive.org/details/{ident}")
                else:
                    console.print("[red][!] No results found[/]")
            except Exception as e:
                console.print(f"[red][!] Error: {e}[/]")

        elif choice == "3":
            banner()
            domain = console.input("[bold yellow]└─[white]$ Enter Domain (e.g. example.com): [/]").strip()
            domain = domain.replace("https://","").replace("http://","").rstrip("/")
            try:
                console.print(f"\n[bold yellow][*] Fetching history for: {domain}[/]\n")
                r = requests.get(
                    f"http://web.archive.org/cdx/search/cdx?url={domain}&output=json&limit=15&fl=timestamp,statuscode,mimetype,length",
                    timeout=15
                ).json()
                table = Table(title=f"[bold yellow]Domain History: {domain}[/]",
                            box=box.DOUBLE_EDGE, border_style="yellow")
                table.add_column("Timestamp", style="yellow", no_wrap=True)
                table.add_column("Status", style="white")
                table.add_column("Type", style="yellow")
                table.add_column("Size", style="white")
                if r and len(r) > 1:
                    for row in r[1:]:
                        if len(row) >= 4:
                            ts = row[0]
                            formatted = f"{ts[:4]}-{ts[4:6]}-{ts[6:8]} {ts[8:10]}:{ts[10:12]}"
                            table.add_row(formatted, str(row[1]), str(row[2]), str(row[3]))
                    console.print(table)
                    console.print(f"\n[bold yellow][*] Full Archive:[/]")
                    print(f"https://web.archive.org/web/*/{domain}")
                else:
                    console.print("[red][!] No history found[/]")
            except Exception as e:
                console.print(f"[red][!] Error: {e}[/]")

        elif choice == "4":
            banner()
            url = console.input("[bold yellow]└─[white]$ Enter URL to Save: [/]").strip()
            if not url.startswith("http"):
                url = "https://" + url
            try:
                console.print(f"\n[bold yellow][*] Saving to Wayback Machine: {url}[/]\n")
                r = requests.get(
                    f"https://web.archive.org/save/{url}",
                    timeout=30,
                    headers={"User-Agent": "Mozilla/5.0"}
                )
                if r.status_code == 200:
                    console.print("[bold green][✓] Page saved successfully![/]")
                    console.print("[bold yellow][*] View at:[/]")
                    print(f"https://web.archive.org/web/*/{url}")
                else:
                    console.print(f"[yellow][!] Status: {r.status_code} — may have been saved[/]")
                    console.print("[bold yellow][*] Check at:[/]")
                    print(f"https://web.archive.org/web/*/{url}")
            except Exception as e:
                console.print(f"[red][!] Error: {e}[/]")
        else:
            console.print("[red][!] Invalid![/]")
        console.input("\n[bold yellow][Press Enter...][/]")

# ══════════════════════════════════════
# MAIN MENU
# ══════════════════════════════════════

def main_menu():
    while True:
        banner()
        console.print(Panel(
            "[bold yellow][1][/]  MetaExtract  — Extract metadata from files\n"
            "[bold yellow][2][/]  MetaEditor   — Edit metadata of files\n"
            "[bold yellow][3][/]  DorkHunter   — Google dorks automation\n"
            "[bold yellow][4][/]  ImgReverse   — Reverse image search\n"
            "[bold yellow][5][/]  TermFetch    — System info\n"
            "[bold yellow][6][/]  ArchiveHunt  — Internet Archive search\n"
            "[bold red][0][/]  Exit",
            title="[bold yellow][ Lemoon v2.0.0 — Main Menu ][/]",
            border_style="yellow"
        ))
        console.print("[bold yellow]┌──[[white]AKAZA SENZO[yellow]㉿[white]Lemoon[yellow]]-[[white]~[yellow]][/]")
        choice = console.input("[bold yellow]└─[white]$ [/]").strip()

        if choice == "1":
            metaextract_menu()
        elif choice == "2":
            metaeditor_menu()
        elif choice == "3":
            dorkhunter_menu()
        elif choice == "4":
            imgreverse_menu()
        elif choice == "5":
            termfetch()
            console.input("\n[bold yellow][Press Enter...][/]")
        elif choice == "6":
            archivehunt_menu()
        elif choice == "0":
            console.print("[bold yellow]\n[!] Exiting Lemoon...[/]")
            sys.exit(0)
        else:
            console.print("[red][!] Invalid![/]")
            console.input("\n[bold yellow][Press Enter...][/]")

if __name__ == "__main__":
    main_menu()
